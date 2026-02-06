from __future__ import annotations

import random
import re
from datetime import datetime
from typing import Dict, List

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.config import RESULTS_DIR


def _format_n(value: float) -> int | float:
    if value is None:
        return 0
    if abs(value - round(value)) < 1e-6:
        return int(round(value))
    return round(value, 1)


def _format_pct(value: float) -> str:
    rounded = round(value, 1)
    if abs(rounded - int(rounded)) < 1e-9:
        return f"{int(rounded)}%"
    text = f"{rounded:.1f}".replace(".", ",")
    return f"{text}%"


def _parse_pct(value: object) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    try:
        text = str(value).strip()
        if not text:
            return None
        return float(text.replace("%", "").replace(",", "."))
    except Exception:
        return None


def _row_value(row: Dict[str, object]) -> float:
    try:
        return float(row.get("n") or 0)
    except Exception:
        return 0.0


def _topic_from_title(title: str) -> str:
    topic = title.strip()
    topic = re.sub(r"\s*\((personas|hogares|viviendas)\)\s*$", "", topic, flags=re.IGNORECASE)
    topic = topic.strip()
    if not topic:
        return "la tabla"
    # Ensure it reads naturally in Spanish
    if topic.lower().startswith(("población", "estado", "materialidad", "servicios", "tenencia", "tic")):
        return f"la {topic.lower()}"
    return topic.lower()


def _clean_label(label: str) -> str:
    text = (label or "").strip()
    text = re.sub(r"^Total\s+de\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^Total\s+", "", text, flags=re.IGNORECASE)
    return text.strip()


def _lower_after_commas(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        return f", {match.group(1).lower()}"

    return re.sub(r",\s*([A-ZÁÉÍÓÚÑ])", repl, text)


def _join_with_y(items: List[str]) -> str:
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} y {items[1]}"
    return f"{', '.join(items[:-1])} y {items[-1]}"


def _build_narrative(rows: List[Dict[str, object]], title: str, denominator_code: str | None = None) -> str:
    source_terms = [
        "De acuerdo al Censo 2024,",
        "Según el Censo 2024,",
        "De acuerdo al Censo (2024),",
    ]
    major_terms = [
        "La mayor proporción corresponde a",
        "La mayoría se concentra en",
        "La categoría predominante es",
        "La proporción más alta se observa en",
    ]
    closing_terms = [
        "A continuación, se presentan los resultados señalados en la siguiente tabla.",
        "A continuación, se presenta la tabla con los resultados indicados.",
        "En la tabla siguiente se detallan los resultados reportados.",
    ]

    base_rows = [
        r
        for r in rows
        if not r.get("is_total")
        and not r.get("is_subtotal")
        and (denominator_code is None or r.get("code") != denominator_code)
    ]
    if not base_rows:
        return f"{random.choice(source_terms)} {random.choice(closing_terms)}"

    topic = _topic_from_title(title)

    if len(base_rows) == 1:
        row = base_rows[0]
        label = _clean_label(row.get("Etiqueta", ""))
        n_val = _row_value(row)
        pct_val = _parse_pct(row.get("Porcentaje"))
        if n_val == 0:
            text = f"{random.choice(source_terms)} respecto de {topic}, no se registran casos en {label} (0 casos)"
        else:
            pct_text = _format_pct(pct_val) if pct_val is not None else ""
            text = f"{random.choice(source_terms)} respecto de {topic}, {label} representa {pct_text} del total observado"
        text = _lower_after_commas(text)
        return f"{text}. {random.choice(closing_terms)}"

    non_zero = [r for r in base_rows if _row_value(r) > 0]
    if not non_zero:
        text = f"{random.choice(source_terms)} respecto de {topic}, no se registran casos con valores distintos de cero"
        text = _lower_after_commas(text)
        return f"{text}. {random.choice(closing_terms)}"
    rows_for_stats = non_zero
    denom = sum(_row_value(r) for r in rows_for_stats) or 0.0

    def row_pct(row: Dict[str, object]) -> float:
        parsed = _parse_pct(row.get("Porcentaje"))
        if parsed is not None:
            return parsed
        if denom <= 0:
            return 0.0
        return (_row_value(row) / denom) * 100

    sorted_rows = sorted(rows_for_stats, key=row_pct, reverse=True)
    leader = sorted_rows[0]
    leader_pct = row_pct(leader)
    leader_label = _clean_label(leader.get("Etiqueta", ""))
    leader_n = _row_value(leader)
    leader_n_text = int(leader_n) if leader_n.is_integer() else leader_n

    parts = [
        f"{random.choice(source_terms)} respecto de {topic}, {random.choice(major_terms)} {leader_label} ({_format_pct(leader_pct)})"
    ]

    if len(sorted_rows) > 1:
        ordered = [f"{_clean_label(r.get('Etiqueta', ''))} ({_format_pct(row_pct(r))})" for r in sorted_rows[1:]]
        if ordered:
            parts.append(f"seguido por {_join_with_y(ordered)}")

    minor_rows = [r for r in sorted_rows if row_pct(r) < 5]
    if len(minor_rows) >= 3:
        minor_sum = sum(row_pct(r) for r in minor_rows)
        parts.append(f"en conjunto, las categorías con menos del 5% suman {_format_pct(minor_sum)}")
    elif minor_rows and len(sorted_rows) > 2:
        smallest = min(minor_rows, key=row_pct)
        parts.append(
            f"por el contrario, la menor presencia se registra en {_clean_label(smallest.get('Etiqueta', ''))} con solo {_format_pct(row_pct(smallest))}"
        )

    text = ", ".join(parts)
    text = _lower_after_commas(text)
    return f"{text}. {random.choice(closing_terms)}"

def _safe_filename(name: str) -> str:
    name = name.strip()
    name = name.replace("/", "_")
    name = re.sub(r"[^A-Za-z0-9_ -]+", "", name)
    name = re.sub(r"\s+", "_", name)
    return name[:80] if name else "grupo"


def _display_col(col: str) -> str:
    return "Frecuencia" if col == "n" else col


def _table_cols(df_rows: pd.DataFrame, category_col: str | None = None) -> tuple[list[str], list[str]]:
    cols = [c for c in df_rows.columns if c not in {"is_total", "is_subtotal", "code"}]
    if category_col and category_col in cols:
        cols = [category_col] + [c for c in cols if c != category_col]
    display_cols = [_display_col(c) for c in cols]
    return cols, display_cols


def _table_df(row_dicts: List[Dict[str, object]], category_col: str | None = None) -> tuple[pd.DataFrame, list[str], list[str], pd.DataFrame]:
    df_rows = pd.DataFrame(row_dicts)
    cols, display_cols = _table_cols(df_rows, category_col)
    df_out = df_rows[cols].copy()
    df_out.columns = display_cols
    return df_rows, cols, display_cols, df_out


def _add_seq_field(paragraph, label: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f"SEQ {label} \\* ARABIC"
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run_result = paragraph.add_run("1")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_result._r.append(fld_end)


def _add_table_caption(doc_ref: Document, title: str, localidad: str) -> None:
    paragraph = doc_ref.add_paragraph()
    try:
        paragraph.style = "Caption"
    except Exception:
        pass
    paragraph.add_run("Tabla ")
    _add_seq_field(paragraph, "Tabla")
    paragraph.add_run(f". {title} - {localidad}")


def _add_source_line(doc_ref: Document) -> None:
    paragraph = doc_ref.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Fuente: Censo, 2024.")
    run.italic = True


def build_reports(
    var_sum: Dict[str, float],
    group_specs: Dict[str, Dict],
    labels: Dict[str, str],
    localidad: str,
    output_prefix: str,
) -> Dict[str, object]:
    reports = []
    loc_slug = _safe_filename(localidad)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    for group_title, spec in group_specs.items():
        vars_list = spec["variables"]
        if not vars_list:
            continue

        category_col = spec.get("category_col")
        category_map = spec.get("category_map", {})
        label_override = spec.get("label_override", {})
        denominator = spec.get("denominator", "sum")
        total_label = spec.get("total_label", "Total")
        no_total = spec.get("no_total", False)

        rows = []

        for code in vars_list:
            value = float(var_sum.get(code, 0.0))
            value_display = _format_n(value)
            row = {
                "Etiqueta": label_override.get(code, labels.get(code, code)),
                "n": value_display,
                "Porcentaje": "",
                "is_total": False,
                "is_subtotal": False,
                "code": code,
            }
            if category_col:
                row[category_col] = category_map.get(code, "")
            rows.append(row)

        if denominator == "by_category" and category_col:
            categories = {}
            for row in rows:
                cat = row.get(category_col, "")
                categories.setdefault(cat, 0.0)
                categories[cat] += float(row["n"]) if row["n"] is not None else 0.0

            for row in rows:
                cat = row.get(category_col, "")
                denom = categories.get(cat, 0.0)
                if denom > 0:
                    row["Porcentaje"] = _format_pct((float(row["n"]) / denom) * 100)

            for cat, subtotal in categories.items():
                subtotal_display = _format_n(subtotal)
                rows.append(
                    {
                        "Etiqueta": total_label,
                        category_col: cat,
                        "n": subtotal_display,
                        "Porcentaje": _format_pct(100) if subtotal > 0 else "",
                        "is_total": True,
                        "is_subtotal": True,
                    }
                )
        else:
            if denominator in var_sum:
                denom_value = float(var_sum.get(denominator, 0.0))
            elif denominator == "sum":
                denom_value = sum(float(r["n"]) for r in rows if r["n"] is not None)
            else:
                denom_value = 0.0

            for row in rows:
                if denom_value > 0 and row["n"] is not None:
                    row["Porcentaje"] = _format_pct((float(row["n"]) / denom_value) * 100)

            if not no_total and total_label:
                total_display = _format_n(denom_value)
                rows.append(
                    {
                        "Etiqueta": total_label,
                        **({category_col: ""} if category_col else {}),
                        "n": total_display,
                        "Porcentaje": _format_pct(100) if denom_value > 0 else "",
                        "is_total": True,
                        "is_subtotal": False,
                    }
                )

        reports.append(
            {
                "title": group_title,
                "rows": rows,
                "csv_path": "",
                "category_col": category_col,
                "denominator": denominator if isinstance(denominator, str) else None,
            }
        )

    # consolidated outputs
    combined_csv = RESULTS_DIR / f"{output_prefix}{loc_slug}_{timestamp}.csv"
    combined_xlsx = RESULTS_DIR / f"{output_prefix}{loc_slug}_{timestamp}.xlsx"
    combined_html = RESULTS_DIR / f"{output_prefix}{loc_slug}_{timestamp}.html"
    combined_docx = RESULTS_DIR / f"{output_prefix}{loc_slug}_{timestamp}.docx"

    all_rows = []
    for rep in reports:
        for row in rep["rows"]:
            out = {"Variable": rep["title"], **{k: v for k, v in row.items() if k not in {"is_total", "is_subtotal"}}}
            all_rows.append(out)

    combined_df = pd.DataFrame(all_rows)
    combined_df_out = combined_df.rename(columns={"n": "Frecuencia"})
    combined_df_out.to_csv(combined_csv, index=False)

    with pd.ExcelWriter(combined_xlsx, engine="openpyxl") as writer:
        combined_df_out.to_excel(writer, index=False, sheet_name="Consolidado")
        used_sheet_names: set[str] = {"Consolidado"}

        def unique_sheet_name(base: str) -> str:
            name = _safe_filename(base)[:31] or "Tabla"
            if name not in used_sheet_names:
                used_sheet_names.add(name)
                return name
            counter = 2
            while True:
                suffix = f"_{counter}"
                trimmed = name[: 31 - len(suffix)]
                candidate = f"{trimmed}{suffix}"
                if candidate not in used_sheet_names:
                    used_sheet_names.add(candidate)
                    return candidate
                counter += 1

        table_entries: List[Dict[str, object]] = []
        for rep in reports:
            table_entries.append(
                {
                    "title": rep["title"],
                    "rows": rep["rows"],
                    "category_col": rep.get("category_col"),
                }
            )
            category_col = rep.get("category_col")
            if category_col:
                categories = [
                    cat
                    for cat in sorted({r.get(category_col, "") for r in rep["rows"]})
                    if cat
                ]
                for cat in categories:
                    sub_rows = [
                        r
                        for r in rep["rows"]
                        if r.get(category_col) == cat
                        or (r.get("is_subtotal") and r.get(category_col) == cat)
                    ]
                    table_entries.append(
                        {
                            "title": f"{rep['title']} - {cat}",
                            "rows": sub_rows,
                            "category_col": rep.get("category_col"),
                        }
                    )

        for idx, entry in enumerate(table_entries, start=1):
            sheet_name = unique_sheet_name(str(entry["title"]))
            _, _, _, df_out = _table_df(entry["rows"], entry.get("category_col"))
            df_out.to_excel(writer, index=False, sheet_name=sheet_name, startrow=1)
            ws = writer.sheets[sheet_name]
            ws.cell(row=1, column=1, value=f"Tabla {idx}. {entry['title']} - {localidad}")

    html_parts = ["<h1>Reporte consolidado</h1>"]
    for rep in reports:
        html_parts.append(f"<h2>{rep['title']}</h2>")
        _, _, _, df_out = _table_df(rep["rows"], rep.get("category_col"))
        html_parts.append(df_out.to_html(index=False))
    combined_html.write_text("\n".join(html_parts), encoding="utf-8")

    try:
        doc = Document()
        doc.add_heading("Reporte consolidado", level=1)

        def add_table(
            doc_ref: Document,
            row_dicts: List[Dict[str, object]],
            category_col: str | None = None,
        ) -> None:
            df_rows, cols, display_cols, _ = _table_df(row_dicts, category_col)
            table = doc_ref.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for idx, col in enumerate(display_cols):
                run = hdr[idx].paragraphs[0].add_run(col)
                run.bold = True
            for _, row in df_rows.iterrows():
                cells = table.add_row().cells
                is_total = bool(row.get("is_total", False))
                for idx, col in enumerate(cols):
                    text = "" if pd.isna(row[col]) else str(row[col])
                    run = cells[idx].paragraphs[0].add_run(text)
                    if is_total:
                        run.bold = True

        # Seccion 1: solo tablas
        doc.add_heading("Sección 1: Tablas", level=1)
        for rep in reports:
            _add_table_caption(doc, rep["title"], localidad)
            add_table(doc, rep["rows"], rep.get("category_col"))
            _add_source_line(doc)

        # Seccion 2: narrativa + tablas
        doc.add_heading("Sección 2: Tablas con narrativa", level=1)
        for rep in reports:
            category_col = rep.get("category_col")
            rows = rep["rows"]
            if category_col:
                categories = [
                    cat
                    for cat in sorted({r.get(category_col, "") for r in rows})
                    if cat
                ]
                for cat in categories:
                    sub_rows = [
                        r
                        for r in rows
                        if r.get(category_col) == cat
                        or (r.get("is_subtotal") and r.get(category_col) == cat)
                    ]
                    title = f"{rep['title']} - {cat}"
                    doc.add_paragraph(
                        _build_narrative(
                            sub_rows,
                            title,
                            rep.get("denominator"),
                        )
                    )
                    _add_table_caption(doc, title, localidad)
                    add_table(doc, sub_rows, rep.get("category_col"))
                    _add_source_line(doc)
            else:
                doc.add_paragraph(_build_narrative(rows, rep["title"], rep.get("denominator")))
                _add_table_caption(doc, rep["title"], localidad)
                add_table(doc, rows, rep.get("category_col"))
                _add_source_line(doc)
        doc.save(combined_docx)
    except Exception:
        combined_docx = RESULTS_DIR / f"{output_prefix}{loc_slug}_{timestamp}_docx_error.txt"
        combined_docx.write_text("Error generando DOCX. Use el HTML o XLSX.")

    return {
        "reports": reports,
        "combined_csv": str(combined_csv),
        "combined_html": str(combined_html),
        "combined_docx": str(combined_docx),
        "combined_xlsx": str(combined_xlsx),
    }
